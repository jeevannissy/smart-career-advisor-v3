# -*- coding: utf-8 -*-
"""
parsing.py -- Document text extraction
Supports PDF (pdfplumber), DOCX (python-docx), and plain text.
"""

import io
import re
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes):
    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2, y_tolerance=3)
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                else:
                    words = page.extract_words()
                    if words:
                        text_parts.append(" ".join(w["text"] for w in words))
    except Exception as e:
        raise ValueError("PDF extraction failed: " + str(e))
    return _clean_text("\n".join(text_parts))


def extract_text_from_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
    except Exception as e:
        raise ValueError("DOCX extraction failed: " + str(e))
    return _clean_text("\n".join(parts))


def extract_text_from_txt(file_bytes):
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return _clean_text(file_bytes.decode(encoding))
        except UnicodeDecodeError:
            continue
    return _clean_text(file_bytes.decode("utf-8", errors="replace"))


def extract_text(file):
    filename = (file.filename or "").lower().strip()
    raw_bytes = file.read()
    if not raw_bytes:
        raise ValueError("Uploaded file is empty.")
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(raw_bytes)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(raw_bytes)
    elif filename.endswith((".txt", ".text")):
        return extract_text_from_txt(raw_bytes)
    else:
        for extractor in (extract_text_from_pdf,
                          extract_text_from_docx,
                          extract_text_from_txt):
            try:
                text = extractor(raw_bytes)
                if text.strip():
                    return text
            except Exception:
                continue
        raise ValueError("Unsupported or unreadable file: " + filename)


def _clean_text(text):
    text = text.replace("\x00", " ")
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()